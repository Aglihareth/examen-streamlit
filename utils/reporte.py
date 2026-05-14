# -*- coding: utf-8 -*-
"""
Created on Wed May 13 11:48:23 2026

@author: calex
"""

from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime

def generar_reporte(nombre, materia, resultado):
    doc = Document()

    doc.add_heading(f'Reporte de Examen — {materia}', 0)
    doc.add_paragraph(f'Alumno: {nombre}')
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'Calificación automática: {resultado["calificacion"]}/{resultado["total"]} pts')
    doc.add_paragraph('')

    doc.add_heading('Detalle de respuestas', level=1)

    for i, detalle in enumerate(resultado['detalles']):
        doc.add_paragraph(f'{i+1}. {detalle["pregunta"]}', style='List Number')
        doc.add_paragraph(f'   Respuesta: {detalle["respuesta_alumno"]}')

        if detalle['tipo'] == 'multiple':
            estado = '✅ Correcto' if detalle['correcto'] else f'❌ Incorrecto (correcta: {detalle["respuesta_correcta"]})'
            doc.add_paragraph(f'   {estado} — {detalle["puntos"]}/{detalle["puntos_posibles"]} pts')
        else:
            doc.add_paragraph(f'   ⏳ Pendiente de revisión — 0/{detalle["puntos_posibles"]} pts')

        doc.add_paragraph('')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()